from docx import Document
import sys
import os
import comtypes.client
document=Document()
document.add_heading('Hello')
document.add_paragraph('My name is xxxx')
document.save('file.docx')
document1=Document()
document1.add_heading('Hii')
document1.add_paragraph('How are u')
document1.save('file1.docx')
document2=Document()
document2.add_paragraph('I am fine')
document2.save('file2.docx')


from docx import Document
files = ['file.docx', 'file1.docx']
wdFormatPDF = 17
combined_document = Document('file2.docx')
count, number_of_files = 0, len(files)
for file in files:
    sub_doc = Document(file)
 
        # Don't add a page break if you've
        # reached the last file.
    if count < number_of_files - 1:
        sub_doc.add_page_break()
 
    for paragraph in sub_doc.paragraphs:
        text = paragraph.text
        combined_document.add_paragraph(text)
    count += 1
 
combined_document.save('file2.docx')
in_file = os.path.abspath('file2.docx')
out_file = os.path.abspath('file2.pdf')
word = comtypes.client.CreateObject('Word.Application')
doc = word.Documents.Open(in_file)
doc.SaveAs(out_file, FileFormat=wdFormatPDF)
doc.Close()
word.Quit()
 

